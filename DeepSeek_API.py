import os
import requests
import json
from docx import Document
from docx.shared import Inches
from docxtpl import DocxTemplate, RichText  # RichText для сложного форматирования, если нужно

# --- Конфигурация ---
CLIENT_DATA_FILE = 'client_data.json'
TEMPLATE_DOCX_FILE = 'bankruptcy_template.docx'
OUTPUT_DOCX_FILE = 'generated_bankruptcy_statement.docx'

# --- DeepSeek API Конфигурация ---
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
DEFAULT_DEEPSEEK_MODEL = "deepseek-chat"
DEEPSEEK_CODER_MODEL = "deepseek-coder"  # Для задач, требующих структурированного вывода (JSON)


# --- 1. Загрузка/Создание данных клиента ---
def get_client_data(filename=CLIENT_DATA_FILE):
    """Загружает данные клиента из JSON файла или возвращает демо-данные."""
    default_data = {
        "fio": "Петров Пётр Петрович",
        "birth_date": "20.07.1978",
        "birth_place": "г. Екатеринбург",
        "address_registration": "г. Екатеринбург, ул. Строителей, д. 10, кв. 5",
        "address_actual": "г. Екатеринбург, ул. Малышева, д. 15, кв. 30",
        "passport_series": "0000",
        "passport_number": "000000",
        "passport_issued_by": "УФМС по Свердловской обл. в Ленинском р-не г. Екатеринбурга",
        "passport_issue_date": "10.08.2001",
        "inn": "000000000000",
        "snils": "000-000-000 00",
        "marital_status": "разведен(а)",
        "children": [],
        "employment_status": "временно не работает",
        "last_work_place": "ИП Сидоров С.С.",
        "last_work_position": "водитель",
        "last_work_dismissal_date": "15.10.2023",
        "income_last_6_months": "50000 руб. (пособие по безработице)",
        "total_debt_amount": 2500000.00,
        "creditors": [
            {
                "name": "Банк 'Восточный Экспресс'",
                "address": "г. Москва, ул. Правды, д. 8",
                "debt_amount": 1200000.00,
                "reason": "Кредитный договор №ZXC/001 от 05.03.2021"
            },
            {
                "name": "МФО 'Быстрые Деньги'",
                "address": "г. Екатеринбург, ул. Вайнера, д. 1",
                "debt_amount": 300000.00,
                "reason": "Договор займа №777 от 10.06.2023"
            },
            {
                "name": "Иванов И.И. (физ. лицо)",
                "address": "г. Екатеринбург, ул. Солнечная, д. 2",
                "debt_amount": 1000000.00,
                "reason": "Расписка от 01.02.2022"
            }
        ],
        "property": [
            {"type": "Доля в квартире (1/2)", "description": "Двухкомнатная, 50 кв.м.",
             "address": "г. Екатеринбург, ул. Строителей, д. 10, кв. 5", "is_pledged": False},
        ],
        "bank_accounts": [
            {"bank_name": "АО 'Уральский Банк Реконструкции и Развития'", "account_number": "40817810100000007777",
             "balance": 530.20}
        ],
        "major_transactions_last_3_years": [
            {"description": "Продажа автомобиля Toyota Camry 2015 г.в.", "date": "01.2023", "amount": 950000.00,
             "counterparty": "Сидоров С.С."}
        ],
        "sro_name": "Ассоциация арбитражных управляющих 'Единство'",
        "sro_address": "620014, г. Екатеринбург, ул. 8 Марта, д. 1, оф. 101",
        "property_notes_from_client": "Еще у меня есть старый ноутбук, но он почти ничего не стоит, и коллекция марок, доставшаяся от деда. Долгов по налогам вроде нет."
    }
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            print(f"Данные клиента загружены из {filename}")
            return json.load(f)
    except FileNotFoundError:
        print(f"Файл {filename} не найден. Используются демонстрационные данные.")
        # Сохраним демо-данные, чтобы пользователь мог их модифицировать
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(default_data, f, ensure_ascii=False, indent=4)
            print(f"Демонстрационные данные сохранены в {filename}")
        except IOError:
            print(f"Не удалось сохранить демонстрационные данные в {filename}.")
        return default_data
    except json.JSONDecodeError:
        print(f"Ошибка декодирования JSON в файле {filename}. Используются демонстрационные данные.")
        return default_data


# --- 2. Функции для работы с DeepSeek API ---
def call_deepseek_api(prompt_messages,
                      model=DEFAULT_DEEPSEEK_MODEL,
                      temperature=0.5,
                      max_tokens=1024,
                      stream=False):
    if not DEEPSEEK_API_KEY:
        print("Критическая ошибка: API ключ DeepSeek (DEEPSEEK_API_KEY) не найден в переменных окружения.")
        return None

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
    }
    payload = {
        "model": model,
        "messages": prompt_messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "stream": stream
    }

    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=90, stream=stream)
        response.raise_for_status()

        if stream:
            full_response_content = ""
            print("Ответ от DeepSeek (потоковый): ", end="")
            for chunk in response.iter_lines():
                if chunk:
                    decoded_chunk = chunk.decode('utf-8')
                    if decoded_chunk.startswith("data: "):
                        json_data_str = decoded_chunk[len("data: "):]
                        if json_data_str.strip() == "[DONE]":
                            break
                        try:
                            data_chunk = json.loads(json_data_str)
                            content_part = data_chunk.get("choices", [{}])[0].get("delta", {}).get("content", "")
                            if content_part:
                                print(content_part, end="", flush=True)
                                full_response_content += content_part
                        except json.JSONDecodeError:
                            # Иногда приходят пустые data: {} или служебные чанки, которые не являются JSON
                            # print(f"\nПропуск чанка (не JSON): {json_data_str}")
                            pass
            print()  # Новая строка после завершения стриминга
            return full_response_content.strip()
        else:
            response_data = response.json()
            if response_data.get("choices") and len(response_data["choices"]) > 0:
                message = response_data["choices"][0].get("message", {})
                return message.get("content", "").strip()
            else:
                print(f"Ошибка: Неожиданный формат ответа от DeepSeek API: {response_data}")
                return None

    except requests.exceptions.RequestException as e:
        print(f"Ошибка при вызове DeepSeek API: {e}")
        if 'response' in locals() and response is not None:  # Проверка, была ли создана переменная response
            print(f"Статус код: {response.status_code}")
            try:
                print(f"Тело ответа: {response.text}")
            except Exception:
                pass
        return None
    except (KeyError, IndexError, TypeError, json.JSONDecodeError) as e:
        print(f"Ошибка при обработке ответа от DeepSeek API: {e}")
        if 'response' in locals() and response is not None and not stream:
            try:
                print(f"Тело ответа: {response.json()}")
            except Exception:
                pass
        return None


def enhance_client_data_with_deepseek(client_data):
    """Обогащает данные клиента, используя DeepSeek."""
    print("\n--- Обогащение данных с помощью DeepSeek API ---")

    # Пример 1: Генерация описания причин банкротства
    reasons_prompt = [
        {"role": "system",
         "content": "Ты юридический ассистент. Твоя задача - на основе предоставленных данных клиента кратко и нейтрально описать возможные причины, приведшие к его неплатежеспособности, для включения в заявление о банкротстве. Сосредоточься на фактах. Формулируй 2-4 предложения."},
        {"role": "user",
         "content": f"Данные клиента:\nСтатус занятости: {client_data.get('employment_status', 'не указан')}\nПоследнее место работы: {client_data.get('last_work_place', 'не указано')}, уволен: {client_data.get('last_work_dismissal_date', 'не указана')}\nОбщая сумма долга: {client_data.get('total_debt_amount', 0):.2f} руб.\nКоличество кредиторов: {len(client_data.get('creditors', []))}.\nДоходы за последние 6 месяцев: {client_data.get('income_last_6_months', 'не указаны')}.\n\nСформулируй краткое описание причин неплатежеспособности."}
    ]
    print("Запрос к DeepSeek для описания причин неплатежеспособности...")
    bankruptcy_reasons = call_deepseek_api(reasons_prompt, temperature=0.3, max_tokens=250, stream=True)
    if bankruptcy_reasons:
        client_data['bankruptcy_reasons_ai_generated'] = bankruptcy_reasons
    else:
        client_data[
            'bankruptcy_reasons_ai_generated'] = "Не удалось автоматически сформулировать причины неплатежеспособности. Требуется ручное заполнение."

    # Пример 2: Извлечение информации из "заметок клиента"
    notes = client_data.get("property_notes_from_client", "")
    if notes:
        property_extraction_prompt = [
            {"role": "system",
             "content": "Ты AI ассистент для извлечения структурированной информации. Из текста ниже извлеки информацию о дополнительном имуществе или обстоятельствах, которые могут быть релевантны для дела о банкротстве. Верни результат в виде краткого списка пунктов или описания. Если ничего релевантного нет, укажи это."},
            {"role": "user", "content": f"Заметки клиента: \"{notes}\"\nИзвлеки релевантную информацию."}
        ]
        print("\nЗапрос к DeepSeek для анализа заметок клиента...")
        extracted_notes_info = call_deepseek_api(property_extraction_prompt, model=DEFAULT_DEEPSEEK_MODEL,
                                                 temperature=0.2, max_tokens=300, stream=False)
        if extracted_notes_info:
            client_data['additional_info_from_notes_ai'] = extracted_notes_info
        else:
            client_data['additional_info_from_notes_ai'] = "Не удалось обработать заметки клиента с помощью ИИ."

    print("--- Обогащение данных завершено ---")
    return client_data


# --- 3. Создание шаблона DOCX (если его нет) ---
def create_template_if_not_exists(filename=TEMPLATE_DOCX_FILE):
    """Создает простой DOCX шаблон, если он не существует."""
    if os.path.exists(filename):
        print(f"Файл шаблона {filename} уже существует.")
        return

    print(f"Создание файла шаблона {filename}...")
    doc = Document()
    # Заголовок
    doc.add_heading('В Арбитражный суд {{ court_name | default("_________________________") }}',
                    level=1)  # Добавил default
    doc.add_paragraph('Адрес: {{ court_address | default("_________________________") }}\n')

    doc.add_paragraph('Должник: {{ fio }}')
    doc.add_paragraph('Дата и место рождения: {{ birth_date }}, {{ birth_place }}')
    doc.add_paragraph('Адрес регистрации: {{ address_registration }}')
    doc.add_paragraph(
        'Адрес фактического проживания: {{ address_actual | default(address_registration) }}')  # Используем actual или registration
    doc.add_paragraph(
        'Паспорт: серия {{ passport_series }} № {{ passport_number }}, выдан {{ passport_issued_by }} {{ passport_issue_date }}')
    doc.add_paragraph('ИНН: {{ inn }}')
    doc.add_paragraph('СНИЛС: {{ snils }}\n')

    doc.add_paragraph('Финансовый управляющий из СРО: {{ sro_name }}')
    doc.add_paragraph('Адрес СРО: {{ sro_address }}\n')

    doc.add_heading('ЗАЯВЛЕНИЕ', level=1)
    doc.add_heading('о признании гражданина банкротом', level=2)
    doc.add_paragraph()

    doc.add_paragraph(
        'Я, {{ fio }}, в соответствии со статьями 213.1 – 213.4 Федерального закона '
        'от 26.10.2002 N 127-ФЗ "О несостоятельности (банкротстве)" (далее – Закон о банкротстве), '
        'прошу признать меня несостоятельным (банкротом).'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Общая сумма моих обязательств перед кредиторами составляет: {{ "%.2f"|format(total_debt_amount) }} руб.')  # Форматирование суммы
    doc.add_paragraph()
    doc.add_paragraph(
        'Обстоятельства, очевидно свидетельствующие о том, что я не в состоянии исполнить денежные обязательства и (или) обязанность по уплате обязательных платежей в установленный срок, и признаки неплатежеспособности и (или) недостаточности имущества у меня имеются.')
    # Вставка сгенерированного текста
    doc.add_paragraph('Краткое описание причин, приведших к неплатежеспособности (сгенерировано ИИ, требует проверки):')
    doc.add_paragraph('{{ bankruptcy_reasons_ai_generated | default("Причины не указаны.") }}\n')

    doc.add_heading('Сведения о кредиторах:', level=3)
    doc.add_paragraph(
        '{% for creditor in creditors %}'
        '- {{ creditor.name }} (адрес: {{ creditor.address | default("не указан") }}), '
        'сумма долга: {{ "%.2f"|format(creditor.debt_amount) }} руб. '
        'Основание: {{ creditor.reason }}\n'
        '{% else %}'
        'Сведения о кредиторах отсутствуют или не указаны.'
        '{% endfor %}\n'
    )

    doc.add_heading('Сведения об имуществе гражданина:', level=3)
    doc.add_paragraph(
        '{% for item in property %}'
        '- {{ item.type }}: {{ item.description }}, адрес: {{ item.address }}. '
        '{% if item.is_pledged %}В залоге у {{ item.pledge_holder | default("не указано") }}.{% endif %}\n'
        '{% else %}'
        'Значимое имущество, подлежащее включению в конкурсную массу, отсутствует (согласно предоставленным данным).'
        '{% endfor %}\n'
    )

    doc.add_heading('Сведения о счетах в банках:', level=3)
    doc.add_paragraph(
        '{% for acc in bank_accounts %}'
        '- {{ acc.bank_name }}, счет № {{ acc.account_number | default("не указан") }}, остаток: {{ "%.2f"|format(acc.balance) }} руб.\n'
        '{% else %}'
        'Счета в банках отсутствуют или не указаны.'
        '{% endfor %}\n'
    )

    # Дополнительная информация из заметок, если есть
    doc.add_paragraph('{% if additional_info_from_notes_ai %}')
    doc.add_heading('Дополнительная информация из заметок клиента (проанализировано ИИ):', level=3)
    doc.add_paragraph('{{ additional_info_from_notes_ai }}')
    doc.add_paragraph('{% endif %}\n')

    doc.add_paragraph('Прошу суд:')
    doc.add_paragraph('1. Признать меня, {{ fio }}, несостоятельным (банкротом).')
    doc.add_paragraph('2. Ввести процедуру реализации имущества гражданина.')
    doc.add_paragraph('3. Утвердить финансовым управляющим члена СРО АУ «{{ sro_name }}».\n')

    doc.add_paragraph('Приложения:')
    doc.add_paragraph('1. Копии документов, подтверждающих обстоятельства (согласно перечню ст. 213.4 Закона).')
    doc.add_paragraph('2. Квитанция об уплате госпошлины (300 руб.).')
    doc.add_paragraph(
        '3. Документ о внесении на депозит суда средств на выплату вознаграждения фин. управляющему (25000 руб.).\n')

    doc.add_paragraph('Дата: «___»___________ 20___ г.         Подпись: _________________ /{{ fio }}/')

    try:
        doc.save(filename)
        print(f"Файл шаблона {filename} успешно создан.")
    except Exception as e:
        print(f"Ошибка при сохранении файла шаблона {filename}: {e}")


# --- 4. Генерация документа ---
def generate_statement_from_template(client_data, template_path, output_path):
    """Генерирует заявление из шаблона DOCX и данных клиента."""
    if not os.path.exists(template_path):
        print(f"Ошибка: Файл шаблона {template_path} не найден. Пожалуйста, создайте его или проверьте путь.")
        return

    tpl = DocxTemplate(template_path)

    # Дополнительные данные, которые могут понадобиться в шаблоне, но нет в client_data
    # Например, имя суда можно запросить у пользователя или взять из настроек
    context = client_data.copy()  # Работаем с копией
    if 'court_name' not in context:
        context['court_name'] = "Арбитражный суд Свердловской области"  # Пример
    if 'court_address' not in context:
        context['court_address'] = "620075, г. Екатеринбург, ул. Шарташская, д. 4"  # Пример

    try:
        tpl.render(context)
        tpl.save(output_path)
        print(f"Заявление успешно сгенерировано и сохранено в: {output_path}")
    except Exception as e:
        print(f"Ошибка при генерации документа из шаблона: {e}")


# --- Основной блок ---
if __name__ == "__main__":
    print("--- Начало процесса генерации заявления о банкротстве ---")

    # 1. Получаем данные клиента
    client_data = get_client_data(CLIENT_DATA_FILE)
    # print("\nИсходные данные клиента:")
    # print(json.dumps(client_data, indent=2, ensure_ascii=False)) # Для отладки

    # 2. Обогащаем данные с помощью DeepSeek API (если ключ доступен)
    if DEEPSEEK_API_KEY:
        client_data = enhance_client_data_with_deepseek(client_data)
        # print("\nОбогащенные данные клиента:")
        # print(json.dumps(client_data, indent=2, ensure_ascii=False)) # Для отладки
    else:
        print("\nAPI ключ DeepSeek не настроен. Пропускаем шаг обогащения данных с помощью ИИ.")
        # Добавим пустые поля, чтобы шаблон не ломался, если ИИ не использовался
        client_data.setdefault('bankruptcy_reasons_ai_generated',
                               'Причины неплатежеспособности не были автоматически сгенерированы.')
        client_data.setdefault('additional_info_from_notes_ai', '')

    # 3. Создаем шаблон DOCX, если его нет
    create_template_if_not_exists(TEMPLATE_DOCX_FILE)

    # 4. Генерируем итоговый документ
    generate_statement_from_template(client_data, TEMPLATE_DOCX_FILE, OUTPUT_DOCX_FILE)

    print("\n--- Процесс завершен ---")
    print(f"Проверьте сгенерированный файл: {os.path.abspath(OUTPUT_DOCX_FILE)}")
    print(f"Данные клиента использовались из: {os.path.abspath(CLIENT_DATA_FILE)}")
    print(f"Шаблон документа: {os.path.abspath(TEMPLATE_DOCX_FILE)}")
    print(
        "\nВАЖНО: Сгенерированный документ является ЧЕРНОВИКОМ и требует обязательной проверки и корректировки юристом!")