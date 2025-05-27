import os
import json
import re  # Для регулярных выражений (NER)

# Библиотеки для PDF
try:
    from pypdf import PdfReader  # Предпочитаемый вариант, если установлен
except ImportError:
    from PyPDF2 import PdfReader  # Альтернатива
from PIL import Image  # Pillow для работы с изображениями
import pytesseract  # Для OCR
import cv2  # OpenCV для предобработки изображений перед OCR

# Библиотеки для Excel и Word
import pandas as pd
from docx import Document as DocxDocument  # Переименовываем, чтобы не конфликтовать с нашим Document из docxtpl

# --- Настройка Tesseract (если необходимо) ---
# Раскомментируйте и укажите ваш путь, если Tesseract не в системном PATH
# if os.name == 'nt': # Для Windows
#     pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
#     # Может потребоваться указать TESSDATA_PREFIX, если языковые файлы не найдены
#     # os.environ['TESSDATA_PREFIX'] = r'C:\Program Files\Tesseract-OCR\tessdata'

# --- Константы и простые паттерны для NER ---
# Эти паттерны очень упрощены. Для продакшена нужны более надежные NER-модели (spaCy, Natasha, etc.)
PATTERNS = {
    "fio": r"([А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)",  # Ф И О
    "birth_date": r"\b(\d{2}\.\d{2}\.\d{4})\b",  # ДД.ММ.ГГГГ
    "passport_series_number": r"\b(\d{2}\s?\d{2})\s*N?\s*(\d{6})\b",  # Серия (XX XX или XXXX) и номер (XXXXXX)
    "inn": r"\bИНН\s*(\d{10}|\d{12})\b",  # ИНН 10 или 12 цифр
    "snils": r"\b(\d{3}-\d{3}-\d{3}\s\d{2})\b"  # СНИЛС XXX-XXX-XXX XX
}


# --- Функции извлечения текста из PDF ---
def preprocess_image_for_ocr(image_path):
    """Предобработка изображения для улучшения качества OCR."""
    try:
        img = cv2.imread(image_path)
        # 1. Преобразование в оттенки серого
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        # 2. Бинаризация (адаптивный порог может быть лучше для некоторых изображений)
        # _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV) # Инвертированный бинарный
        thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
        # 3. Удаление шума (опционально)
        # denoised = cv2.medianBlur(thresh, 3)
        # cv2.imwrite("temp_processed_ocr_page.png", thresh) # Для отладки
        return Image.fromarray(thresh)  # pytesseract работает с PIL Image
    except Exception as e:
        print(f"Ошибка при предобработке изображения {image_path}: {e}")
        # В случае ошибки, пытаемся прочитать как есть
        return Image.open(image_path)


def extract_text_from_pdf(pdf_path):
    """Извлекает текст из PDF. Сначала пытается извлечь текстовый слой, потом OCR."""
    text_content = ""
    try:
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
            # Попытка 1: Извлечь текстовый слой
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                try:
                    text_content += page.extract_text() or ""
                except Exception as e_extract:
                    print(f"Ошибка при извлечении текста со страницы {page_num + 1} (текстовый слой): {e_extract}")

            if text_content.strip():  # Если текст извлечен из текстового слоя
                print(f"Текст извлечен из текстового слоя PDF: {pdf_path}")
                return text_content

            # Попытка 2: OCR, если текстовый слой пуст или его нет
            print(f"Текстовый слой в PDF {pdf_path} пуст или отсутствует. Попытка OCR...")
            text_content_ocr = ""
            # Конвертируем PDF в изображения (требует pdf2image или poppler-utils)
            # Для простоты здесь будем считать, что PDF - это уже изображение или
            # мы используем внешнюю утилиту для конвертации страниц PDF в PNG/JPG.
            # Вместо этого, можно использовать PdfReader для извлечения изображений со страницы,
            # если они там есть, или рендерить страницы в изображения.

            # Упрощенный вариант: считаем, что если текстового слоя нет, то весь PDF - это скан
            # и пытаемся применить OCR ко всему файлу (Tesseract может это делать для некоторых PDF)
            # Более надежно - конвертировать каждую страницу в изображение и применять OCR.
            # Для примера, мы будем использовать только первую страницу PDF для OCR, если он многостраничный.
            # Иначе код усложнится (нужна библиотека типа pdf2image)

            # Эта часть с Tesseract напрямую для PDF может не всегда хорошо работать.
            # Лучше конвертировать страницы в PNG/JPG.
            try:
                # Tesseract может пытаться обработать PDF напрямую, если он содержит изображения
                text_content_ocr = pytesseract.image_to_string(pdf_path, lang='rus+eng')
                if text_content_ocr.strip():
                    print(f"Текст извлечен из PDF через OCR (прямая обработка): {pdf_path}")
                    return text_content_ocr
            except pytesseract.TesseractError as e_tess_direct:
                print(f"Ошибка Tesseract при прямой обработке PDF {pdf_path}: {e_tess_direct}")
            except Exception as e_direct_ocr:
                print(f"Общая ошибка при OCR PDF напрямую {pdf_path}: {e_direct_ocr}")

            # Если выше не сработало, и если у нас есть способ конвертировать PDF в изображения:
            # Допустим, у нас есть папка с изображениями страниц PDF (page_1.png, page_2.png ...)
            # Или мы используем библиотеку pdf2image для этого:
            # from pdf2image import convert_from_path
            # try:
            #     images = convert_from_path(pdf_path, dpi=300) # dpi для качества
            #     for i, image_pil in enumerate(images):
            #         # Сохраняем временное изображение для OpenCV или используем PIL напрямую
            #         temp_image_path = f"temp_page_{i}.png"
            #         image_pil.save(temp_image_path, "PNG")
            #         processed_image = preprocess_image_for_ocr(temp_image_path)
            #         page_text = pytesseract.image_to_string(processed_image, lang='rus+eng')
            #         text_content_ocr += page_text + "\n"
            #         os.remove(temp_image_path) # Удаляем временный файл
            #     if text_content_ocr.strip():
            #         print(f"Текст извлечен из PDF через OCR (страница за страницей): {pdf_path}")
            #         return text_content_ocr
            # except Exception as e_pdf2image:
            #     print(f"Ошибка при конвертации PDF в изображения для OCR ({pdf_path}): {e_pdf2image}. Убедитесь, что poppler установлен.")
            #     return "Ошибка OCR: не удалось обработать PDF как изображение."

            # Если ни один из методов OCR для PDF не сработал и нет текстового слоя
            if not text_content.strip() and not text_content_ocr.strip():
                return "Не удалось извлечь текст из PDF (ни текстовый слой, ни OCR)."


    except Exception as e:
        print(f"Критическая ошибка при обработке PDF {pdf_path}: {e}")
        return f"Ошибка чтения PDF: {e}"
    return text_content if text_content.strip() else "Текст не извлечен из PDF."


# --- Функция извлечения текста из изображений (для OCR сканов не-PDF) ---
def extract_text_from_image(image_path):
    """Извлекает текст из файла изображения с помощью OCR."""
    try:
        processed_image = preprocess_image_for_ocr(image_path)
        text = pytesseract.image_to_string(processed_image, lang='rus+eng')  # 'rus+eng' для русского и английского
        print(f"Текст извлечен из изображения через OCR: {image_path}")
        return text
    except pytesseract.TesseractNotFoundError:
        print("Ошибка: Tesseract OCR не найден. Установите его и/или укажите путь в pytesseract.tesseract_cmd.")
        return "Ошибка OCR: Tesseract не найден."
    except Exception as e:
        print(f"Ошибка при OCR изображения {image_path}: {e}")
        return f"Ошибка OCR: {e}"


# --- Функция извлечения данных из Excel ---
def extract_data_from_excel(excel_path):
    """Извлекает данные из Excel файла. Возвращает словарь, где ключ - имя листа, значение - DataFrame в виде списка словарей."""
    try:
        xls = pd.ExcelFile(excel_path)
        data = {}
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            # Преобразуем NaN в None или пустые строки для лучшей сериализации в JSON
            df = df.fillna("")  # или None, если предпочитаете
            data[sheet_name] = df.to_dict(orient='records')
        print(f"Данные извлечены из Excel: {excel_path}")
        return data
    except Exception as e:
        print(f"Ошибка при чтении Excel файла {excel_path}: {e}")
        return {"error": f"Ошибка чтения Excel: {e}"}


# --- Функция извлечения текста из Word ---
def extract_text_from_word(word_path):
    """Извлекает текст из Word (.docx) файла."""
    try:
        doc = DocxDocument(word_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Также можно извлекать текст из таблиц, если это необходимо
        # for table in doc.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             full_text.append(cell.text)
        content = '\n'.join(full_text)
        print(f"Текст извлечен из Word: {word_path}")
        return content
    except Exception as e:
        print(f"Ошибка при чтении Word файла {word_path}: {e}")
        return f"Ошибка чтения Word: {e}"


# --- Простой NER для извлечения структурированных данных из текста ---
def simple_ner_from_text(text_content):
    """Извлекает базовые сущности из текста с помощью регулярных выражений."""
    extracted_entities = {}
    for entity_name, pattern in PATTERNS.items():
        matches = re.findall(pattern, text_content)
        if matches:
            if entity_name == "passport_series_number" and isinstance(matches[0], tuple):
                extracted_entities["passport_series"] = matches[0][0].replace(" ", "")
                extracted_entities["passport_number"] = matches[0][1]
            elif len(matches) == 1:
                extracted_entities[entity_name] = matches[0]
            else:
                extracted_entities[entity_name] = matches  # Список, если несколько совпадений
    return extracted_entities


# --- Диспетчер обработки файлов ---
def process_document(file_path):
    """
    Определяет тип файла и вызывает соответствующую функцию для извлечения данных.
    Возвращает словарь с извлеченными данными.
    """
    _, file_extension = os.path.splitext(file_path.lower())
    filename = os.path.basename(file_path)
    print(f"\nОбработка файла: {filename} (тип: {file_extension})")

    extracted_data = {"filename": filename, "file_type": file_extension, "content": None, "structured_data": {}}

    if not os.path.exists(file_path):
        extracted_data["error"] = "Файл не найден."
        print(f"Ошибка: Файл {file_path} не найден.")
        return extracted_data

    try:
        if file_extension == '.pdf':
            text = extract_text_from_pdf(file_path)
            extracted_data["content"] = text
            if not text.startswith("Ошибка") and not text.startswith("Не удалось"):
                extracted_data["structured_data"] = simple_ner_from_text(text)
        elif file_extension in ['.xls', '.xlsx']:
            excel_content = extract_data_from_excel(file_path)
            extracted_data["content"] = excel_content  # Здесь content - это структурированные данные
            # Если это анкета, то можно попытаться преобразовать в единую структуру
            # Например, если ожидается лист 'Анкета Клиента'
            if isinstance(excel_content, dict) and "Анкета Клиента" in excel_content:
                # Предполагаем, что анкета - это одна строка или простой формат ключ-значение
                # Для простоты берем первую строку, если есть данные
                if excel_content["Анкета Клиента"]:
                    # Это очень упрощенно. Реальный маппинг будет сложнее.
                    extracted_data["structured_data"] = excel_content["Анкета Клиента"][0]
            elif isinstance(excel_content, dict) and "error" not in excel_content:
                # Если нет конкретного листа, но есть данные, можно попробовать их объединить
                # или просто оставить как есть
                pass  # Оставляем в extracted_data["content"]
            else:
                extracted_data["error"] = excel_content.get("error", "Не удалось извлечь данные из Excel")


        elif file_extension == '.docx':
            text = extract_text_from_word(file_path)
            extracted_data["content"] = text
            if not text.startswith("Ошибка"):
                extracted_data["structured_data"] = simple_ner_from_text(text)
        elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            text = extract_text_from_image(file_path)
            extracted_data["content"] = text
            if not text.startswith("Ошибка"):
                extracted_data["structured_data"] = simple_ner_from_text(text)
        else:
            extracted_data["error"] = f"Неподдерживаемый тип файла: {file_extension}"
            print(f"Файл {filename} имеет неподдерживаемый тип: {file_extension}")
    except Exception as e:
        extracted_data["error"] = f"Общая ошибка при обработке файла {filename}: {e}"
        print(f"Критическая ошибка при обработке файла {filename}: {e}")

    return extracted_data


# --- Основной блок для демонстрации ---
if __name__ == "__main__":
    # Создадим несколько тестовых файлов для демонстрации
    # (в реальном сценарии эти файлы будут предоставляться пользователем)

    # 1. Тестовый DOCX
    try:
        doc_test = DocxDocument()
        doc_test.add_paragraph("Анкета клиента")
        doc_test.add_paragraph("ФИО: Сидоров Сидор Сидорович")
        doc_test.add_paragraph("Дата рождения: 10.12.1990")
        doc_test.add_paragraph("Паспорт: 45 05 123456, выдан МВД гор. Примерный 01.01.2010")
        doc_test.add_paragraph("ИНН 1234567890")  # без префикса "ИНН "
        doc_test.add_paragraph("СНИЛС: 111-222-333 44")
        doc_test.add_paragraph("Прочая информация о доходах и расходах.")
        test_docx_file = "test_document.docx"
        doc_test.save(test_docx_file)
        print(f"Создан тестовый файл: {test_docx_file}")
    except Exception as e:
        print(f"Не удалось создать тестовый DOCX: {e}")
        test_docx_file = None

    # 2. Тестовый Excel
    try:
        excel_data = {
            'ФИО': ['Иванов Иван Иванович'],
            'Дата рождения': ['15.05.1985'],
            'ИНН': ['123456789012'],
            'СНИЛС': ['123-456-789 00']
        }
        df_test = pd.DataFrame(excel_data)
        test_excel_file = "test_data.xlsx"
        df_test.to_excel(test_excel_file, sheet_name="Анкета Клиента", index=False)
        print(f"Создан тестовый файл: {test_excel_file}")
    except Exception as e:
        print(f"Не удалось создать тестовый Excel: {e}")
        test_excel_file = None

    # 3. Тестовый PDF (текстовый) - создадим его из DOCX для простоты
    # Это требует наличия LibreOffice или MS Word для конвертации.
    # Пропустим этот шаг автоматического создания для упрощения примера.
    # Вместо этого, предположим, что у вас есть PDF файл, например, 'sample.pdf'
    # Для теста создайте простой PDF вручную или используйте существующий.
    # Я создам "заглушку" для PDF - текстовый файл, чтобы Tesseract мог его обработать как "плохой" PDF
    test_pdf_file_text_layer = "sample_text.pdf"  # Пользователь должен предоставить этот файл
    test_pdf_file_scan = "sample_scan.pdf"  # Пользователь должен предоставить этот файл (изображение как PDF)
    test_image_file = "sample_scan.png"  # Пользователь должен предоставить этот файл

    # Если вы хотите протестировать PDF, поместите файлы с такими именами в папку со скриптом.
    # Например, создайте sample_text.pdf с текстовым слоем (сохраните Word/LibreOffice документ как PDF).
    # И sample_scan.pdf (скан документа, сохраненный как PDF).
    # И sample_scan.png (скан документа как изображение).

    # --- Пример использования ---
    documents_to_process = []
    if test_docx_file and os.path.exists(test_docx_file):
        documents_to_process.append(test_docx_file)
    if test_excel_file and os.path.exists(test_excel_file):
        documents_to_process.append(test_excel_file)

    # Добавьте пути к вашим реальным PDF и изображениям для теста
    if os.path.exists("sample_text.pdf"): documents_to_process.append("sample_text.pdf")
    if os.path.exists("sample_scan.pdf"): documents_to_process.append("sample_scan.pdf")
    if os.path.exists("sample_scan.png"): documents_to_process.append("sample_scan.png")
    if not documents_to_process:
        print(
            "\nНе найдены тестовые файлы для обработки. Поместите sample_text.pdf, sample_scan.pdf или sample_scan.png в папку.")

    all_extracted_data = []
    for doc_path in documents_to_process:
        data = process_document(doc_path)
        all_extracted_data.append(data)
        print("-" * 30)

    print("\n--- Итоги извлечения ---")
    # Выведем структурированные данные, если они есть
    for item in all_extracted_data:
        print(f"\nФайл: {item['filename']}")
        if item.get("error"):
            print(f"  Ошибка: {item['error']}")
        else:
            print(f"  Тип: {item['file_type']}")
            if item['file_type'] == '.xlsx':  # Для Excel контент уже структурирован
                print(f"  Содержимое (Excel):")
                for sheet_name, records in item.get("content", {}).items():
                    if sheet_name == 'error': continue
                    print(f"    Лист '{sheet_name}':")
                    for i, record in enumerate(records[:2]):  # Показать первые 2 записи для краткости
                        print(f"      Запись {i + 1}: {record}")
                    if len(records) > 2: print("      ...")
            else:  # Для PDF, Word, изображений
                print(f"  Извлеченный текст (первые 200 симв.):\n'{str(item.get('content', ''))[:200]}...'")

            if item.get("structured_data"):
                print(f"  Извлеченные сущности (NER):")
                for key, value in item["structured_data"].items():
                    print(f"    {key}: {value}")
            else:
                print("  Структурированные сущности не извлечены (или не применялся NER).")

    # Дальше эти `all_extracted_data` можно агрегировать и передавать
    # в модуль структурирования и нормализации, а затем в DeepSeek и генератор документов.